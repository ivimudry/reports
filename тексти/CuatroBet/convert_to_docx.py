"""
Convert task-spec .md files to .docx for the client.
Handles: headings, tables, bold, lists, horizontal rules.
"""
import re, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = r"c:\Projects\REPORTS\тексти\CuatroBet\task-specs"
DST = r"c:\Projects\REPORTS\тексти\CuatroBet\task-specs-docx"
os.makedirs(DST, exist_ok=True)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(sh)

def add_formatted_text(paragraph, text):
    """Parse bold (**text**) and code (`text`) in a line and add runs."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        else:
            paragraph.add_run(part)

def process_md(filepath):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 3))
            add_formatted_text(h, heading_text)
            i += 1
            continue

        # Table block
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            if len(table_lines) < 2:
                for tl in table_lines:
                    doc.add_paragraph(tl)
                continue

            # Parse header
            header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            # Skip separator line
            data_start = 1
            if re.match(r'^[\s|:\-]+$', table_lines[1]):
                data_start = 2

            data_rows = []
            for tl in table_lines[data_start:]:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                data_rows.append(cells)

            cols = len(header_cells)
            if cols == 0:
                continue

            table = doc.add_table(rows=1 + len(data_rows), cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header row
            for j, hc in enumerate(header_cells):
                cell = table.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(hc.replace('**', ''))
                run.bold = True
                run.font.size = Pt(9)
                set_cell_shading(cell, 'E8E8F0')

            # Data rows
            for ri, row_data in enumerate(data_rows):
                for j in range(cols):
                    cell = table.rows[ri + 1].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    text = row_data[j] if j < len(row_data) else ''
                    add_formatted_text(p, text)
                    for run in p.runs:
                        run.font.size = Pt(9)

            doc.add_paragraph('')  # spacer
            continue

        # Blockquote
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # List items
        m_list = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if m_list:
            indent = len(m_list.group(1))
            text = m_list.group(3)
            bullet = m_list.group(2)
            if bullet in ['-', '*']:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph(style='List Number')
            if indent > 2:
                p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, text)
            i += 1
            continue

        # Checkbox items
        m_cb = re.match(r'^- \[([ x])\]\s+(.*)', line)
        if m_cb:
            checked = m_cb.group(1) == 'x'
            text = m_cb.group(2)
            prefix = '☑ ' if checked else '☐ '
            p = doc.add_paragraph()
            add_formatted_text(p, prefix + text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1

    return doc


# Process all .md files
for fname in sorted(os.listdir(SRC)):
    if not fname.endswith('.md'):
        continue
    src_path = os.path.join(SRC, fname)
    dst_name = fname.replace('.md', '.docx')
    dst_path = os.path.join(DST, dst_name)
    print(f"Converting {fname} -> {dst_name} ... ", end='')
    try:
        doc = process_md(src_path)
        doc.save(dst_path)
        print("OK")
    except Exception as e:
        print(f"ERROR: {e}")

print("\nDone! All files in:", DST)
