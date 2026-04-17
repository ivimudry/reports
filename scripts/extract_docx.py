import docx
import sys

doc = docx.Document(r'тексти\CuatroBet\CuatroBet_Argentina_плейбук.docx')

# Print paragraphs
for i, p in enumerate(doc.paragraphs):
    txt = p.text[:200] if p.text else '(empty)'
    print(f'[{i}] style={p.style.name} | {txt}')

# Print tables
for ti, table in enumerate(doc.tables):
    print(f'\n=== TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ===')
    for ri, row in enumerate(table.rows):
        cells = [cell.text[:80] for cell in row.cells]
        print(f'  Row {ri}: {cells}')
