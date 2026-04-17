import docx

doc = docx.Document(r'тексти\CuatroBet\CuatroBet_Argentina_плейбук.docx')

output = []

# Paragraphs
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    style = p.style.name
    if style == 'Heading 1':
        output.append(f'\n# {txt}\n')
    elif style == 'Heading 2':
        output.append(f'\n## {txt}\n')
    elif style == 'Heading 3':
        output.append(f'\n### {txt}\n')
    elif style == 'List Paragraph' or style.startswith('List'):
        output.append(f'- {txt}')
    else:
        output.append(txt)

# Tables
for ti, table in enumerate(doc.tables):
    output.append(f'\n<!-- TABLE {ti} -->')
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    output.append('| ' + ' | '.join(headers) + ' |')
    output.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
    for row in table.rows[1:]:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        output.append('| ' + ' | '.join(cells) + ' |')
    output.append('')

with open(r'тексти\CuatroBet\cuatrobet_playbook_en.md', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print(f'Done. {len(output)} lines written.')
