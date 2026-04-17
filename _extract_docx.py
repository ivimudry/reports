import docx, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document(r'c:\Projects\REPORTS\тексти\CuatroBet\CuatroBet_Argentina_плейбук.docx')

for p in doc.paragraphs:
    style = p.style.name if p.style else ''
    text = p.text.strip()
    if not text:
        print()
        continue
    if 'Heading 1' in style:
        print(f'# {text}')
    elif 'Heading 2' in style:
        print(f'## {text}')
    elif 'Heading 3' in style:
        print(f'### {text}')
    elif 'Heading 4' in style:
        print(f'#### {text}')
    elif p.style and 'List' in style:
        print(f'- {text}')
    else:
        print(text)

# Also extract tables
for i, table in enumerate(doc.tables):
    print(f'\n---TABLE {i+1}---')
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(' | '.join(cells))
